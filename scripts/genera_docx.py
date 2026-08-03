"""
Genera i documenti Word a partire dai Markdown di `docs/`.

Perche' esiste
--------------
La documentazione si scrive in Markdown: si legge su GitHub, git ne mostra le
differenze riga per riga e due persone possono modificarla senza conflitti
insanabili. I file Word servono per la consegna e per la stampa.

Tenere allineate a mano due versioni dello stesso documento non funziona: prima
o poi divergono. Qui il `.md` e' l'unica fonte di verita' e il `.docx` e' un
prodotto, rigenerabile in qualsiasi momento.

Uso
---
    python scripts/genera_docx.py            # tutti i documenti di docs/
    python scripts/genera_docx.py docs/M2-dati/dati.md
    python scripts/genera_docx.py --verifica  # controlla se i .docx sono aggiornati

Il Markdown riconosciuto e' quello effettivamente usato nel progetto: titoli,
paragrafi, elenchi puntati e numerati, tabelle, blocchi di codice, citazioni,
righe orizzontali, e in linea grassetto, corsivo, codice e collegamenti.
"""

from __future__ import annotations

import argparse
import re
import sys
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

RADICE = Path(__file__).resolve().parent.parent
CARTELLA_DOCS = RADICE / "docs"

# --------------------------------------------------------------------------- #
# Identita' visiva
# --------------------------------------------------------------------------- #

AMBRA = RGBColor(0xA8, 0x72, 0x00)  # accento del progetto, scuro quanto basta
NAVY = RGBColor(0x1A, 0x21, 0x38)  # titoli e testo forte
GRIGIO = RGBColor(0x5B, 0x61, 0x78)  # sottotitoli e note
BRUNO = RGBColor(0x7A, 0x52, 0x00)  # codice in linea: leggibile senza gridare

# I documenti usano `#` per il titolo e `##` per le sezioni principali: la
# gerarchia visiva parte quindi dal livello 2.
LIVELLI = {
    1: (15, NAVY),
    2: (13, NAVY),
    3: (11.5, AMBRA),
    4: (10.5, GRIGIO),
}
BORDO = "D9D2C4"
FONDO_INTESTAZIONE = "F5EFE1"
FONDO_CODICE = "F4F4F2"

FONT_TESTO = "Calibri"
FONT_CODICE = "Consolas"

OCCHIELLO = "BEEWATCH AI  ·  PROGETTO 7"

# Data fissa scritta nelle proprieta' del documento. Serve a rendere la
# generazione deterministica: con la data corrente, ogni esecuzione produrrebbe
# byte diversi anche a contenuto identico, e git segnalerebbe come modificati
# tutti i .docx a ogni rigenerazione.
DATA_FISSA = datetime(2026, 1, 1)  # noqa: DTZ001 - proprieta' del documento, non un istante
DATA_ZIP = (2026, 1, 1, 0, 0, 0)

# Nome esteso delle milestone, ricavato dalla cartella che contiene il file.
MILESTONE = {
    "M1": "Fondamenta e setup",
    "M2": "Dati e database",
    "M3": "Accesso ai dati ed ETL",
    "M4": "Machine Learning",
    "M5": "Componente AI generativa",
    "M6": "Interfaccia Streamlit",
    "M7": "Qualita, Docker, etica",
    "M8": "Documentazione, demo e consegna",
    "MX": "Estensioni opzionali",
}


# --------------------------------------------------------------------------- #
# Utilita' di basso livello (python-docx non le espone)
# --------------------------------------------------------------------------- #


def _sfondo_paragrafo(paragrafo, colore: str) -> None:
    ombreggiatura = OxmlElement("w:shd")
    ombreggiatura.set(qn("w:val"), "clear")
    ombreggiatura.set(qn("w:fill"), colore)
    paragrafo._p.get_or_add_pPr().append(ombreggiatura)


def _sfondo_cella(cella, colore: str) -> None:
    ombreggiatura = OxmlElement("w:shd")
    ombreggiatura.set(qn("w:val"), "clear")
    ombreggiatura.set(qn("w:fill"), colore)
    cella._tc.get_or_add_tcPr().append(ombreggiatura)


def _bordo(paragrafo, lato: str, colore: str, spessore: int = 6) -> None:
    """Aggiunge un bordo a un lato del paragrafo (usato per righe e citazioni)."""
    pPr = paragrafo._p.get_or_add_pPr()
    bordi = pPr.find(qn("w:pBdr"))
    if bordi is None:
        bordi = OxmlElement("w:pBdr")
        pPr.append(bordi)
    elemento = OxmlElement(f"w:{lato}")
    elemento.set(qn("w:val"), "single")
    elemento.set(qn("w:sz"), str(spessore))
    elemento.set(qn("w:space"), "6")
    elemento.set(qn("w:color"), colore)
    bordi.append(elemento)


LARGHEZZA_UTILE_CM = 21.59 - 3.2 * 2  # pagina meno i margini
LARGHEZZA_MINIMA_CM = 1.8


def _distribuisci_colonne(tabella, celle: list[list[str]]) -> None:
    """Assegna a ogni colonna una larghezza proporzionale al suo contenuto.

    Word e LibreOffice interpretano l'adattamento automatico in modo diverso e
    il risultato e' imprevedibile. Calcolarle qui rende il documento identico
    ovunque venga aperto.
    """
    numero_colonne = len(celle[0])
    pesi = []
    for indice in range(numero_colonne):
        testi = [riga[indice] for riga in celle if indice < len(riga)]
        if not testi:
            pesi.append(1.0)
            continue
        # La radice attenua le colonne con una sola cella molto lunga; la parola
        # piu' lunga fa da pavimento, perche' non si puo' mandare a capo.
        piu_lunga = max((len(p) for t in testi for p in t.split()), default=1)
        pesi.append(max(max(len(t) for t in testi) ** 0.5, piu_lunga * 0.85))

    totale = sum(pesi)
    disponibile = LARGHEZZA_UTILE_CM - LARGHEZZA_MINIMA_CM * numero_colonne
    for indice, colonna in enumerate(tabella.columns):
        larghezza = Cm(LARGHEZZA_MINIMA_CM + disponibile * pesi[indice] / totale)
        colonna.width = larghezza
        for cella in colonna.cells:
            cella.width = larghezza


def _bordi_tabella(tabella, colore: str = BORDO) -> None:
    proprieta = tabella._tbl.tblPr
    bordi = OxmlElement("w:tblBorders")
    for lato in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elemento = OxmlElement(f"w:{lato}")
        elemento.set(qn("w:val"), "single")
        elemento.set(qn("w:sz"), "4")
        elemento.set(qn("w:space"), "0")
        elemento.set(qn("w:color"), colore)
        bordi.append(elemento)
    proprieta.append(bordi)


# --------------------------------------------------------------------------- #
# Testo in linea
# --------------------------------------------------------------------------- #

_INLINE = re.compile(
    r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*|\[[^\]]+\]\([^)]+\))",
    re.S,
)


def scrivi_testo(paragrafo, testo: str, *, dimensione: float = 10.5, colore=None) -> None:
    """Scrive il testo dentro il paragrafo interpretando la formattazione in linea."""
    for pezzo in _INLINE.split(testo):
        if not pezzo:
            continue
        grassetto = corsivo = codice = False
        if pezzo.startswith("**") and pezzo.endswith("**"):
            pezzo, grassetto = pezzo[2:-2], True
        elif pezzo.startswith("`") and pezzo.endswith("`"):
            pezzo, codice = pezzo[1:-1], True
        elif pezzo.startswith("*") and pezzo.endswith("*"):
            pezzo, corsivo = pezzo[1:-1], True
        elif collegamento := re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", pezzo):
            etichetta, indirizzo = collegamento.groups()
            pezzo = etichetta if indirizzo.startswith("#") else f"{etichetta} ({indirizzo})"

        run = paragrafo.add_run(pezzo)
        run.bold = grassetto
        run.italic = corsivo
        run.font.size = Pt(dimensione - 0.5 if codice else dimensione)
        run.font.name = FONT_CODICE if codice else FONT_TESTO
        run.font.color.rgb = BRUNO if codice else (colore or NAVY)


# --------------------------------------------------------------------------- #
# Costruzione del documento
# --------------------------------------------------------------------------- #


@dataclass
class Intestazione:
    """I tre elementi in cima a ogni documento."""

    occhiello: str
    titolo: str
    sottotitolo: str


def prepara_documento() -> Document:
    documento = Document()

    proprieta = documento.core_properties
    proprieta.author = "BeeWatch AI"
    proprieta.last_modified_by = "scripts/genera_docx.py"
    proprieta.created = proprieta.modified = DATA_FISSA
    proprieta.revision = 1

    sezione = documento.sections[0]
    sezione.page_width, sezione.page_height = Cm(21.59), Cm(27.94)
    sezione.top_margin = sezione.bottom_margin = Cm(2.5)
    sezione.left_margin = sezione.right_margin = Cm(3.2)

    normale = documento.styles["Normal"]
    normale.font.name = FONT_TESTO
    normale.font.size = Pt(10.5)
    normale.font.color.rgb = NAVY
    normale.paragraph_format.space_after = Pt(8)
    normale.paragraph_format.line_spacing = 1.15

    for livello, (dimensione, colore) in LIVELLI.items():
        prima = {1: 20, 2: 18, 3: 12, 4: 10}[livello]
        stile = documento.styles[f"Heading {livello}"]
        stile.font.name = FONT_TESTO
        stile.font.size = Pt(dimensione)
        stile.font.bold = True
        stile.font.color.rgb = colore
        stile.paragraph_format.space_before = Pt(prima)
        stile.paragraph_format.space_after = Pt(4)
        stile.paragraph_format.keep_with_next = True

    return documento


def scrivi_intestazione(documento: Document, testa: Intestazione) -> None:
    paragrafo = documento.add_paragraph()
    paragrafo.paragraph_format.space_after = Pt(2)
    run = paragrafo.add_run(testa.occhiello)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = FONT_TESTO
    run.font.color.rgb = AMBRA

    paragrafo = documento.add_paragraph()
    paragrafo.paragraph_format.space_after = Pt(2)
    run = paragrafo.add_run(testa.titolo)
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = FONT_TESTO
    run.font.color.rgb = NAVY

    paragrafo = documento.add_paragraph()
    paragrafo.paragraph_format.space_after = Pt(16)
    _bordo(paragrafo, "bottom", BORDO.replace("#", ""))
    if testa.sottotitolo:
        run = paragrafo.add_run(testa.sottotitolo)
        run.font.size = Pt(10.5)
        run.font.name = FONT_TESTO
        run.font.color.rgb = GRIGIO


def scrivi_pie_di_pagina(documento: Document, etichetta: str) -> None:
    piede = documento.sections[0].footer
    paragrafo = piede.paragraphs[0]
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragrafo.add_run(etichetta)
    run.font.size = Pt(8)
    run.font.name = FONT_TESTO
    run.font.color.rgb = GRIGIO


# --------------------------------------------------------------------------- #
# Conversione del Markdown
# --------------------------------------------------------------------------- #


def _riga_tabella(riga: str) -> bool:
    return riga.startswith("|") and riga.endswith("|")


def _separatore_tabella(riga: str) -> bool:
    return bool(re.fullmatch(r"\|[\s:|-]+\|", riga))


def aggiungi_tabella(documento: Document, righe: list[str]) -> None:
    """Costruisce una tabella Word dalle righe Markdown."""
    celle = [[c.strip() for c in r.strip("|").split("|")] for r in righe]
    intestazione, corpo = celle[0], celle[2:]  # celle[1] e' il separatore

    # Alcune tabelle del progetto sono usate come elenchi a due colonne e hanno
    # l'intestazione vuota: in quel caso la riga non va disegnata affatto.
    con_intestazione = any(testo for testo in intestazione)

    tabella = documento.add_table(rows=1 if con_intestazione else 0, cols=len(intestazione))
    tabella.alignment = WD_TABLE_ALIGNMENT.LEFT
    tabella.autofit = False
    _bordi_tabella(tabella)

    if con_intestazione:
        for cella, testo in zip(tabella.rows[0].cells, intestazione, strict=False):
            _sfondo_cella(cella, FONDO_INTESTAZIONE)
            paragrafo = cella.paragraphs[0]
            paragrafo.paragraph_format.space_after = Pt(2)
            run = paragrafo.add_run(testo)
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.name = FONT_TESTO
            run.font.color.rgb = NAVY

    for valori in corpo:
        riga = tabella.add_row()
        for cella, testo in zip(riga.cells, valori, strict=False):
            paragrafo = cella.paragraphs[0]
            paragrafo.paragraph_format.space_after = Pt(2)
            scrivi_testo(paragrafo, testo, dimensione=9.5)

    _distribuisci_colonne(tabella, celle[:1] + celle[2:] if con_intestazione else celle[2:])
    documento.add_paragraph().paragraph_format.space_after = Pt(6)


def aggiungi_codice(documento: Document, righe: list[str]) -> None:
    for riga in righe:
        paragrafo = documento.add_paragraph()
        paragrafo.paragraph_format.space_after = Pt(0)
        paragrafo.paragraph_format.left_indent = Cm(0.4)
        paragrafo.paragraph_format.line_spacing = 1.0
        _sfondo_paragrafo(paragrafo, FONDO_CODICE)
        run = paragrafo.add_run(riga if riga.strip() else " ")
        run.font.name = FONT_CODICE
        run.font.size = Pt(9)
        run.font.color.rgb = NAVY
    documento.add_paragraph().paragraph_format.space_after = Pt(6)


def converti(percorso_md: Path, percorso_docx: Path) -> None:
    """Legge un Markdown e scrive il Word corrispondente."""
    righe = percorso_md.read_text(encoding="utf-8").splitlines()
    documento = prepara_documento()

    indice = 0
    titolo = percorso_md.stem
    if righe and righe[0].startswith("# "):
        titolo = righe[0][2:].strip()
        indice = 1

    codice_milestone = percorso_md.parent.name.split("-")[0]
    nome_milestone = MILESTONE.get(codice_milestone, "")
    sottotitolo = f"{codice_milestone} — {nome_milestone}" if nome_milestone else ""

    scrivi_intestazione(documento, Intestazione(OCCHIELLO, titolo, sottotitolo))
    scrivi_pie_di_pagina(documento, f"BeeWatch AI  ·  {percorso_md.name}")

    dentro_codice = False
    blocco_codice: list[str] = []
    paragrafo_corrente: list[str] = []

    def chiudi_paragrafo() -> None:
        if not paragrafo_corrente:
            return
        paragrafo = documento.add_paragraph()
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        scrivi_testo(paragrafo, " ".join(paragrafo_corrente))
        paragrafo_corrente.clear()

    while indice < len(righe):
        riga = righe[indice].rstrip()
        nuda = riga.strip()

        if nuda.startswith("```"):
            if dentro_codice:
                aggiungi_codice(documento, blocco_codice)
                blocco_codice = []
            else:
                chiudi_paragrafo()
            dentro_codice = not dentro_codice
            indice += 1
            continue

        if dentro_codice:
            blocco_codice.append(riga)
            indice += 1
            continue

        if not nuda:
            chiudi_paragrafo()
            indice += 1
            continue

        if _riga_tabella(nuda):
            chiudi_paragrafo()
            blocco: list[str] = []
            while indice < len(righe) and _riga_tabella(righe[indice].strip()):
                blocco.append(righe[indice].strip())
                indice += 1
            if len(blocco) >= 2 and _separatore_tabella(blocco[1]):
                aggiungi_tabella(documento, blocco)
            continue

        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", nuda):
            chiudi_paragrafo()
            paragrafo = documento.add_paragraph()
            paragrafo.paragraph_format.space_before = Pt(6)
            paragrafo.paragraph_format.space_after = Pt(10)
            _bordo(paragrafo, "bottom", BORDO)
            indice += 1
            continue

        titolo_trovato = re.match(r"(#{1,4})\s+(.*)", nuda)
        if titolo_trovato:
            chiudi_paragrafo()
            livello = min(len(titolo_trovato.group(1)), 4)
            dimensione, colore = LIVELLI[livello]
            paragrafo = documento.add_paragraph(style=f"Heading {livello}")
            scrivi_testo(paragrafo, titolo_trovato.group(2), dimensione=dimensione, colore=colore)
            indice += 1
            continue

        if nuda.startswith(">"):
            chiudi_paragrafo()
            citazione: list[str] = []
            while indice < len(righe) and righe[indice].strip().startswith(">"):
                citazione.append(righe[indice].strip().lstrip(">").strip())
                indice += 1
            paragrafo = documento.add_paragraph()
            paragrafo.paragraph_format.left_indent = Cm(0.5)
            paragrafo.paragraph_format.space_after = Pt(10)
            _bordo(paragrafo, "left", "A87200", spessore=12)
            testo_citazione = " ".join(r for r in citazione if r)
            scrivi_testo(paragrafo, testo_citazione, dimensione=10, colore=GRIGIO)
            continue

        elenco = re.match(r"[-*]\s+(.*)", nuda)
        if elenco:
            chiudi_paragrafo()
            paragrafo = documento.add_paragraph(style="List Bullet")
            paragrafo.paragraph_format.space_after = Pt(3)
            scrivi_testo(paragrafo, elenco.group(1))
            indice += 1
            continue

        numerato = re.match(r"\d+\.\s+(.*)", nuda)
        if numerato:
            chiudi_paragrafo()
            paragrafo = documento.add_paragraph(style="List Number")
            paragrafo.paragraph_format.space_after = Pt(3)
            scrivi_testo(paragrafo, numerato.group(1))
            indice += 1
            continue

        paragrafo_corrente.append(nuda)
        indice += 1

    chiudi_paragrafo()
    percorso_docx.parent.mkdir(parents=True, exist_ok=True)
    documento.save(percorso_docx)
    _normalizza_archivio(percorso_docx)


def _normalizza_archivio(percorso: Path) -> None:
    """Riscrive il .docx con date e ordine fissi dentro l'archivio.

    Un file .docx e' uno zip, e zipfile ci scrive dentro l'ora di creazione di
    ogni voce. Senza questo passaggio due esecuzioni a contenuto identico
    produrrebbero byte diversi, e git segnalerebbe come modificati tutti i
    documenti a ogni rigenerazione.
    """
    with zipfile.ZipFile(percorso) as archivio:
        contenuti = {voce.filename: archivio.read(voce.filename) for voce in archivio.infolist()}

    # `[Content_Types].xml` resta in testa come vuole il formato OPC.
    ordine = sorted(contenuti, key=lambda n: (n != "[Content_Types].xml", n))

    with zipfile.ZipFile(percorso, "w", zipfile.ZIP_DEFLATED) as archivio:
        for nome in ordine:
            informazioni = zipfile.ZipInfo(nome, date_time=DATA_ZIP)
            informazioni.compress_type = zipfile.ZIP_DEFLATED
            informazioni.external_attr = 0o600 << 16
            archivio.writestr(informazioni, contenuti[nome])


# --------------------------------------------------------------------------- #
# Interfaccia a riga di comando
# --------------------------------------------------------------------------- #


def elenca_markdown(percorsi: list[str]) -> list[Path]:
    if percorsi:
        return [Path(p).resolve() for p in percorsi]
    return sorted(p for p in CARTELLA_DOCS.rglob("*.md") if p.name != "README.md")


def main() -> int:
    lettore = argparse.ArgumentParser(description="Genera i .docx dai .md di docs/")
    lettore.add_argument("file", nargs="*", help="Markdown da convertire (vuoto = tutti)")
    lettore.add_argument(
        "--verifica",
        action="store_true",
        help="non scrive nulla: segnala i .docx mancanti o piu' vecchi del .md",
    )
    argomenti = lettore.parse_args()

    documenti = elenca_markdown(argomenti.file)
    if not documenti:
        print("Nessun documento da convertire in docs/")
        return 0

    obsoleti = []
    for percorso_md in documenti:
        percorso_docx = percorso_md.with_suffix(".docx")
        relativo = percorso_md.relative_to(RADICE)

        if argomenti.verifica:
            if not percorso_docx.exists():
                obsoleti.append(f"{relativo}  ->  .docx mancante")
            elif percorso_docx.stat().st_mtime < percorso_md.stat().st_mtime:
                obsoleti.append(f"{relativo}  ->  .docx piu' vecchio del .md")
            continue

        converti(percorso_md, percorso_docx)
        print(f"  {relativo}  ->  {percorso_docx.name}")

    if argomenti.verifica:
        if obsoleti:
            print("Documenti Word non aggiornati:")
            for voce in obsoleti:
                print(" ", voce)
            return 1
        print(f"Tutti i {len(documenti)} documenti Word sono aggiornati.")
        return 0

    print(f"\nGenerati {len(documenti)} documenti Word.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
